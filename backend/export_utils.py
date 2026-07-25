"""Shared Excel / PDF / Word export builders for tabular reports."""
from __future__ import annotations

import io
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional, Sequence

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter

LOGO_PATH = Path(__file__).parent / "assets" / "gti-logo.png"

_MIME = {
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def export_mime(fmt: str) -> str:
    key = (fmt or "").lower().strip()
    if key not in _MIME:
        raise ValueError(f"Unsupported export format: {fmt}")
    return _MIME[key]


def export_filename(base: str, fmt: str) -> str:
    key = (fmt or "").lower().strip()
    if key not in _MIME:
        raise ValueError(f"Unsupported export format: {fmt}")
    return f"{base}.{key}"


def build_export_bytes(
    title: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    fmt: str,
) -> bytes:
    key = (fmt or "").lower().strip()
    if key == "xlsx":
        return _build_xlsx(title, headers, rows)
    if key == "pdf":
        return _build_pdf(title, headers, rows)
    if key == "docx":
        return _build_docx(title, headers, rows)
    raise ValueError(f"Unsupported export format: {fmt}")


def _cell_str(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    return str(value)


def _build_xlsx(
    title: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = (title or "Export")[:31]

    header_fill = PatternFill("solid", fgColor="1E3A5F")
    header_font = Font(bold=True, color="FFFFFF")

    for col, name in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=name)
        cell.fill = header_fill
        cell.font = header_font

    for row_idx, row in enumerate(rows, start=2):
        for col, value in enumerate(row, start=1):
            ws.cell(row=row_idx, column=col, value=_cell_str(value))

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 18

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _build_docx(
    title: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta = doc.add_paragraph(f"Generated: {generated}")
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    if LOGO_PATH.is_file():
        try:
            doc.add_picture(str(LOGO_PATH), width=Inches(1.2))
        except Exception:
            pass

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, name in enumerate(headers):
        hdr_cells[i].text = name
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            cells[c_idx].text = _cell_str(value)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(
    title: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    page_width, page_height = landscape(A4)
    left_margin = 0.45 * inch
    right_margin = 0.45 * inch
    usable_width = page_width - left_margin - right_margin

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=landscape(A4),
        leftMargin=left_margin,
        rightMargin=right_margin,
        topMargin=0.4 * inch,
        bottomMargin=0.4 * inch,
    )
    base = getSampleStyleSheet()
    header_style = ParagraphStyle(
        "ExportHeader",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=8,
        textColor=colors.white,
        leading=10,
        alignment=TA_LEFT,
    )
    cell_style = ParagraphStyle(
        "ExportCell",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=7.5,
        leading=9.5,
        alignment=TA_LEFT,
        wordWrap="CJK",
    )
    title_style = ParagraphStyle(
        "ExportTitle",
        parent=base["Heading1"],
        fontSize=16,
        leading=20,
        spaceAfter=2,
    )
    brand_style = ParagraphStyle(
        "ExportBrand",
        parent=base["Heading2"],
        fontSize=11,
        leading=14,
        spaceAfter=2,
    )
    meta_style = ParagraphStyle(
        "ExportMeta",
        parent=base["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#64748B"),
        spaceAfter=8,
    )

    story = []

    logo_flowable = _pdf_logo_flowable(max_width=1.1 * inch, max_height=0.55 * inch)
    if logo_flowable is not None:
        story.append(logo_flowable)
        story.append(Spacer(1, 6))

    story.append(Paragraph("GTI Fleet Solutions", brand_style))
    story.append(Paragraph(_escape_pdf_text(title), title_style))
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    story.append(Paragraph(f"Generated: {generated}", meta_style))

    col_widths = _pdf_column_widths(headers, usable_width)
    data = [[Paragraph(_escape_pdf_text(h), header_style) for h in headers]]
    for row in rows:
        data.append([
            Paragraph(_escape_pdf_text(_cell_str(v)), cell_style)
            for v in row
        ])

    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    doc.build(story)
    return buf.getvalue()


def _escape_pdf_text(value: object) -> str:
    text = _cell_str(value)
    text = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    # Allow wrapping of long unbroken tokens (UUIDs, regs, etc.)
    return _insert_soft_breaks(text)


def _insert_soft_breaks(text: str, chunk: int = 12) -> str:
    """Insert zero-width spaces so long tokens wrap inside PDF table cells."""
    parts = []
    for token in text.split(" "):
        if len(token) <= chunk:
            parts.append(token)
            continue
        pieces = [token[i:i + chunk] for i in range(0, len(token), chunk)]
        parts.append("&#8203;".join(pieces))
    return " ".join(parts)


def _pdf_column_widths(headers: Sequence[str], usable_width: float) -> List[float]:
    """Prefer wider columns for text-heavy fields; keep date/number cols compact."""
    weights = []
    for header in headers:
        key = str(header).strip().lower()
        if key in {"date", "fuel (l)", "fuel", "distance (km)", "distance", "overall status", "status"}:
            weights.append(1.0)
        elif key in {"vehicle", "driver"}:
            weights.append(1.6)
        elif key in {"route", "purpose", "description"}:
            weights.append(2.4)
        else:
            weights.append(1.4)
    total = sum(weights) or 1.0
    return [usable_width * (w / total) for w in weights]


def _pdf_logo_flowable(max_width: float, max_height: float):
    """Load and scale the GTI logo so a tall asset cannot dominate the page."""
    if not LOGO_PATH.is_file():
        return None
    try:
        from reportlab.platypus import Image
        from reportlab.lib.utils import ImageReader

        reader = ImageReader(str(LOGO_PATH))
        iw, ih = reader.getSize()
        if not iw or not ih:
            return None
        scale = min(max_width / float(iw), max_height / float(ih))
        width = float(iw) * scale
        height = float(ih) * scale
        img = Image(str(LOGO_PATH), width=width, height=height)
        img.hAlign = "LEFT"
        return img
    except Exception:
        return None
